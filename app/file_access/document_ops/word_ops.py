# app/file_access/document_ops/word_ops.py
"""
Protocol-agnostic Word document operations.

All functions take a FileStorageProvider and operate on bytes,
making them work with any storage backend (local, SMB, NFS, etc.).
"""
import io
from typing import Any, Dict, List, Optional
from datetime import datetime

import structlog

from app.file_access.base_fs import FileStorageProvider

# Word library imports
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger = structlog.get_logger()
    logger.warning("python_docx_not_available", message="Install python-docx for Word operations: pip install python-docx")

logger = structlog.get_logger()


async def read_word(
    provider: FileStorageProvider,
    path: str
) -> Document:
    """
    Read Word document from any provider.
    
    Args:
        provider: FileStorageProvider instance (connected)
        path: Path to Word document (.docx)
        
    Returns:
        python-docx Document object
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ImportError: If python-docx is not installed
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word operations. Install: pip install python-docx")
    
    try:
        logger.debug("reading_word", path=path)
        
        # Read document bytes from provider
        data = await provider.read_file(path)
        
        # Load document from bytes
        buffer = io.BytesIO(data)
        doc = Document(buffer)
        
        logger.info("word_read_success", path=path, paragraphs=len(doc.paragraphs))
        return doc
        
    except Exception as e:
        logger.error("word_read_failed", path=path, error=str(e))
        raise


async def write_word(
    provider: FileStorageProvider,
    path: str,
    document: Document,
    overwrite: bool = False
) -> bool:
    """
    Write Word document to any provider.
    
    Args:
        provider: FileStorageProvider instance (connected)
        path: Destination path
        document: python-docx Document to write
        overwrite: Whether to overwrite existing file
        
    Returns:
        True if successful
        
    Raises:
        FileExistsError: If file exists and overwrite=False
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word operations. Install: pip install python-docx")
    
    try:
        logger.debug("writing_word", path=path, overwrite=overwrite)
        
        # Save document to bytes
        buffer = io.BytesIO()
        document.save(buffer)
        data = buffer.getvalue()
        
        # Write bytes to provider
        result = await provider.write_file(path, data, overwrite=overwrite)
        
        if result.success:
            logger.info("word_write_success", path=path, size=len(data))
            return True
        else:
            logger.error("word_write_failed", path=path, error=result.error)
            raise Exception(f"Failed to write Word document: {result.error}")
            
    except Exception as e:
        logger.error("word_write_failed", path=path, error=str(e))
        raise


async def create_word(
    provider: FileStorageProvider,
    path: str,
    content: Dict[str, Any],
    overwrite: bool = False
) -> bool:
    """
    Create Word document with structured content.
    
    Args:
        provider: FileStorageProvider instance (connected)
        path: Destination path
        content: Dictionary with document structure
        overwrite: Whether to overwrite existing file
        
    Returns:
        True if successful
        
    Content structure:
    {
        "title": "Document Title",
        "sections": [
            {
                "heading": "Section 1",
                "paragraphs": ["Paragraph 1 text", "Paragraph 2 text"]
            },
            {
                "heading": "Section 2",
                "paragraphs": ["More content"],
                "table": {
                    "headers": ["Column 1", "Column 2"],
                    "rows": [
                        ["Value 1", "Value 2"],
                        ["Value 3", "Value 4"]
                    ]
                }
            }
        ]
    }
    
    Example:
        content = {
            "title": "Quote #Q001",
            "sections": [
                {
                    "heading": "Customer Information",
                    "paragraphs": [
                        "Customer: ACME Corp",
                        "Contact: John Doe",
                        "Email: john@acme.com"
                    ]
                },
                {
                    "heading": "Quote Items",
                    "table": {
                        "headers": ["Item", "Quantity", "Price"],
                        "rows": [
                            ["Product A", "2", "$500.00"],
                            ["Product B", "1", "$300.00"]
                        ]
                    }
                }
            ]
        }
        
        await create_word(provider, "/quotes/Q001.docx", content)
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word operations. Install: pip install python-docx")
    
    try:
        logger.debug("creating_word", path=path)
        
        # Create new document
        doc = Document()
        
        # Add title
        if "title" in content:
            title = doc.add_heading(content["title"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sections
        for section in content.get("sections", []):
            # Add section heading
            if "heading" in section:
                doc.add_heading(section["heading"], level=1)
            
            # Add paragraphs
            for para_text in section.get("paragraphs", []):
                doc.add_paragraph(para_text)
            
            # Add table if present
            if "table" in section:
                table_data = section["table"]
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                
                # Create table
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"
                
                # Add headers
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = str(header)
                
                # Add rows
                for row_idx, row_data in enumerate(rows, start=1):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_value in enumerate(row_data):
                        row_cells[col_idx].text = str(cell_value)
            
            # Add space after section
            doc.add_paragraph()
        
        # Write document
        success = await write_word(provider, path, doc, overwrite=overwrite)
        
        logger.info("word_create_success", path=path, sections=len(content.get("sections", [])))
        return success
        
    except Exception as e:
        logger.error("word_create_failed", path=path, error=str(e))
        raise


async def update_word(
    provider: FileStorageProvider,
    path: str,
    replacements: Dict[str, str]
) -> bool:
    """
    Update Word document by replacing placeholder text.
    
    Args:
        provider: FileStorageProvider instance (connected)
        path: Path to Word document
        replacements: Dict of {placeholder: replacement_text}
        
    Returns:
        True if successful
        
    Example:
        replacements = {
            "{{CUSTOMER_NAME}}": "ACME Corp",
            "{{QUOTE_NUMBER}}": "Q001",
            "{{AMOUNT}}": "$1,000.00",
            "{{DATE}}": "2024-01-15"
        }
        
        await update_word(provider, "/templates/quote_template.docx", replacements)
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word operations. Install: pip install python-docx")
    
    try:
        logger.debug("updating_word", path=path, replacements_count=len(replacements))
        
        # Read document
        doc = await read_word(provider, path)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Write back
        success = await write_word(provider, path, doc, overwrite=True)
        
        logger.info("word_update_success", path=path, replacements=list(replacements.keys()))
        return success
        
    except Exception as e:
        logger.error("word_update_failed", path=path, error=str(e))
        raise


async def extract_word_text(
    provider: FileStorageProvider,
    path: str
) -> str:
    """
    Extract plain text from Word document.
    
    Args:
        provider: FileStorageProvider instance (connected)
        path: Path to Word document
        
    Returns:
        Extracted text content
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word operations. Install: pip install python-docx")
    
    try:
        logger.debug("extracting_word_text", path=path)
        
        # Read document
        doc = await read_word(provider, path)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(p.text for p in cell.paragraphs if p.text.strip())
                    row_text.append(cell_text)
                text_parts.append(" | ".join(row_text))
        
        full_text = "\n".join(text_parts)
        
        logger.info("word_text_extracted", path=path, length=len(full_text))
        return full_text
        
    except Exception as e:
        logger.error("word_text_extraction_failed", path=path, error=str(e))
        raise


async def add_word_paragraph(
    provider: FileStorageProvider,
    path: str,
    text: str,
    style: Optional[str] = None
) -> bool:
    """
    Append paragraph to existing Word document.
    
    Args:
        provider: FileStorageProvider instance (connected)
        path: Path to Word document
        text: Paragraph text to add
        style: Paragraph style name (optional)
        
    Returns:
        True if successful
        
    Example:
        await add_word_paragraph(
            provider,
            "/documents/report.docx",
            "This is a new paragraph added to the document."
        )
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word operations. Install: pip install python-docx")
    
    try:
        logger.debug("adding_word_paragraph", path=path)
        
        # Read document
        doc = await read_word(provider, path)
        
        # Add paragraph
        paragraph = doc.add_paragraph(text)
        if style:
            paragraph.style = style
        
        # Write back
        success = await write_word(provider, path, doc, overwrite=True)
        
        logger.info("word_paragraph_added", path=path)
        return success
        
    except Exception as e:
        logger.error("word_paragraph_add_failed", path=path, error=str(e))
        raise


async def add_word_table(
    provider: FileStorageProvider,
    path: str,
    headers: List[str],
    rows: List[List[str]]
) -> bool:
    """
    Append table to existing Word document.
    
    Args:
        provider: FileStorageProvider instance (connected)
        path: Path to Word document
        headers: Column headers
        rows: Table rows
        
    Returns:
        True if successful
        
    Example:
        await add_word_table(
            provider,
            "/documents/report.docx",
            headers=["Product", "Quantity", "Price"],
            rows=[
                ["Product A", "10", "$100.00"],
                ["Product B", "5", "$50.00"]
            ]
        )
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word operations. Install: pip install python-docx")
    
    try:
        logger.debug("adding_word_table", path=path, rows=len(rows))
        
        # Read document
        doc = await read_word(provider, path)
        
        # Add table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
        
        # Add rows
        for row_idx, row_data in enumerate(rows, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = str(cell_value)
        
        # Write back
        success = await write_word(provider, path, doc, overwrite=True)
        
        logger.info("word_table_added", path=path, rows=len(rows))
        return success
        
    except Exception as e:
        logger.error("word_table_add_failed", path=path, error=str(e))
        raise
